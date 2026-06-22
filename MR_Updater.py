#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Medical Report Updater
----------------------
- Updates the date/time in the header, footers and signature to "now"
  (signature keeps a +5 second gap, exactly like the original files).
- Optionally updates the "Created on:" vascular-access date to today.
- For the "Clinical Chemistry" and "Hematology" lab tables:
      * inserts a NEW left-most column with the lab date you enter,
      * pushes the other columns one step to the right,
      * drops the old right-most (oldest) column,
      * fills the new column with a copy of the old oldest column's values.
- Fonts / formatting / layout of the original file are preserved.

Author: built for Ahmed
"""

import os
import re
import copy
from datetime import datetime, timedelta

from docx import Document
from docx.oxml.ns import qn

# ----------------------------------------------------------------------------
# Date helpers
# ----------------------------------------------------------------------------

# Matches "Apr-05-2026 8:57:54 AM" (3-letter month, time WITH seconds + AM/PM).
# Clinical dates such as "May-11-2015" (no time) are NOT matched -> stay safe.
DATETIME_RE = re.compile(r'[A-Za-z]{3}-\d{2}-\d{4}\s+\d{1,2}:\d{2}:\d{2}\s+[AP]M')


def fmt_datetime(dt):
    """Format like the report header: 'Apr-05-2026 8:57:54 AM' (hour NOT zero-padded)."""
    hour12 = dt.strftime('%I').lstrip('0') or '12'
    return dt.strftime('%b-%d-%Y ') + hour12 + dt.strftime(':%M:%S %p')


def fmt_labdate(dt):
    """Format like the lab columns: 'M/D/YYYY' (no zero padding)."""
    return f"{dt.month}/{dt.day}/{dt.year}"


def parse_labdate(text):
    """Accept M/D/YYYY, MM/DD/YYYY, D-M-YYYY or YYYY-MM-DD; return a datetime."""
    text = text.strip()
    for fmt in ("%m/%d/%Y", "%d/%m/%Y", "%Y-%m-%d", "%m-%d-%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue
    raise ValueError("Could not understand the lab date. Use M/D/YYYY, e.g. 4/15/2026")


# ----------------------------------------------------------------------------
# Low-level docx helpers (formatting-preserving)
# ----------------------------------------------------------------------------

def iter_paragraphs(container):
    """Yield every paragraph in a header/footer/body, including those inside tables."""
    for p in container.paragraphs:
        yield p
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def replace_in_paragraph(paragraph, pattern, replacement):
    """
    Replace every regex match inside a paragraph WITHOUT disturbing the
    formatting of the rest of the paragraph. The matched text is rewritten
    using the formatting of the first run it touches.
    Returns the number of replacements made.
    """
    runs = paragraph.runs
    if not runs:
        return 0
    full = ''.join(r.text for r in runs)
    matches = list(pattern.finditer(full))      # found ONCE on original text
    if not matches:
        return 0
    # Apply right-to-left so positions of earlier matches stay valid
    # (and the replacement text is never re-scanned).
    for m in reversed(matches):
        ms, me = m.start(), m.end()
        runs = paragraph.runs
        spans = []
        pos = 0
        for i, r in enumerate(runs):
            spans.append((i, pos, pos + len(r.text)))
            pos += len(r.text)
        affected = [(i, s, e) for (i, s, e) in spans if e > ms and s < me]
        if not affected:
            continue
        first_i = affected[0][0]
        for i, s, e in affected:
            run = runs[i]
            local = run.text
            keep_before = local[:max(0, ms - s)]
            keep_after = local[max(0, me - s):]
            run.text = (keep_before + replacement + keep_after) if i == first_i \
                else (keep_before + keep_after)
    return len(matches)


def _clone_rpr(template_run, dst_run):
    """Copy run-properties (font, size, bold, ...) from template_run to dst_run."""
    if template_run is None:
        return
    src_rpr = template_run._element.find(qn('w:rPr'))
    if src_rpr is None:
        return
    dst_el = dst_run._element
    existing = dst_el.find(qn('w:rPr'))
    if existing is not None:
        dst_el.remove(existing)
    dst_el.insert(0, copy.deepcopy(src_rpr))


def set_cell_text(cell, text, template_run=None):
    """Set a cell's text while keeping its existing run formatting.
    Handles multi-line values (joined by '\\n') with proper line breaks and
    never leaves a stray empty trailing paragraph."""
    lines = text.split('\n')
    paras = cell.paragraphs
    p = paras[0]
    # remove any extra paragraphs entirely (no empty trailing lines)
    for extra in paras[1:]:
        extra._p.getparent().remove(extra._p)
    # keep the first run (its formatting), drop the rest
    if p.runs:
        first = p.runs[0]
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        first = p.add_run('')
        _clone_rpr(template_run, first)
    first.text = lines[0]
    for line in lines[1:]:
        first.add_break()
        first.add_text(line)


def _find_template_run(table):
    """Find a representative formatted run from a data cell, for empty-cell formatting."""
    for ri, row in enumerate(table.rows):
        if ri == 0:
            continue
        for ci, cell in enumerate(row.cells):
            if ci == 0:
                continue
            for p in cell.paragraphs:
                if p.runs and p.runs[0].text.strip():
                    return p.runs[0]
    return None


# ----------------------------------------------------------------------------
# The three edit operations
# ----------------------------------------------------------------------------

def update_datetimes(doc, now_dt, signature_gap_seconds=5):
    """Header + footers -> now; body signature -> now + gap. Returns count."""
    now_str = fmt_datetime(now_dt)
    sig_str = fmt_datetime(now_dt + timedelta(seconds=signature_gap_seconds))
    n = 0

    for section in doc.sections:
        for attr in ('header', 'even_page_header', 'first_page_header',
                     'footer', 'even_page_footer', 'first_page_footer'):
            hf = getattr(section, attr, None)
            if hf is None:
                continue
            for p in iter_paragraphs(hf):
                n += replace_in_paragraph(p, DATETIME_RE, now_str)

    # Body (the signature is the only datetime here) -> now + 5s
    body_container = doc  # Document exposes .paragraphs and .tables
    for p in iter_paragraphs(body_container):
        n += replace_in_paragraph(p, DATETIME_RE, sig_str)

    return n


def roll_lab_table(table, new_date_str):
    """
    Insert new_date_str as the new left-most (newest) column, push the rest
    right by one, drop the oldest (right-most) column. The new column's data
    cells are filled with a copy of the previous oldest column's values.
    Table layout: col0 = label, cols 1..6 = date columns (1 newest .. 6 oldest).
    """
    template_run = _find_template_run(table)
    for ri, row in enumerate(table.rows):
        cells = row.cells
        if len(cells) < 7:
            continue
        vals = [cells[c].text for c in range(1, 7)]   # current cols 1..6
        if ri == 0:
            new_first = new_date_str                  # header gets the user's date
        else:
            new_first = vals[5]                       # data gets a copy of the oldest
        new_order = [new_first, vals[0], vals[1], vals[2], vals[3], vals[4]]
        for idx, c in enumerate(range(1, 7)):
            set_cell_text(cells[c], new_order[idx], template_run)


def update_lab_tables(doc, lab_date_dt, table_names=("Clinical Chemistry", "Hematology")):
    """Roll the requested lab tables forward with the given lab date."""
    new_date_str = fmt_labdate(lab_date_dt)
    done = []
    for t in doc.tables:
        head = t.rows[0].cells[0].text.strip()
        if head in table_names and len(t.rows[0].cells) >= 7:
            roll_lab_table(t, new_date_str)
            done.append(head)
    return done


# ----------------------------------------------------------------------------
# Orchestration
# ----------------------------------------------------------------------------

def process_report(input_path, lab_date_text, output_path=None, now_dt=None):
    """
    Run all edits and save a NEW file. Returns (output_path, summary_dict).
    'Created on' and 'First use on' dates are never modified.
    """
    if now_dt is None:
        now_dt = datetime.now()
    lab_dt = parse_labdate(lab_date_text)

    doc = Document(input_path)

    dt_count = update_datetimes(doc, now_dt)
    labs_done = update_lab_tables(doc, lab_dt)

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        stamp = now_dt.strftime("%Y-%m-%d")
        output_path = f"{base}_updated_{stamp}{ext}"

    doc.save(output_path)

    summary = {
        "datetimes_updated": dt_count,
        "lab_tables_rolled": labs_done,
        "now": fmt_datetime(now_dt),
        "signature": fmt_datetime(now_dt + timedelta(seconds=5)),
        "lab_date": fmt_labdate(lab_dt),
    }
    return output_path, summary


# ----------------------------------------------------------------------------
# GUI
# ----------------------------------------------------------------------------

def launch_gui():
    import customtkinter as ctk
    from tkinter import filedialog, messagebox

    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("blue")

    app = ctk.CTk()
    app.title("Medical Report Updater")
    app.geometry("640x560")
    app.minsize(560, 520)

    state = {"path": None}

    # --- Title ---
    ctk.CTkLabel(app, text="Medical Report Updater",
                 font=ctk.CTkFont(size=22, weight="bold")).pack(pady=(20, 4))
    ctk.CTkLabel(app, text="Updates header/footer/signature date & rolls the lab tables forward.",
                 font=ctk.CTkFont(size=12), text_color="gray70").pack(pady=(0, 16))

    frame = ctk.CTkFrame(app)
    frame.pack(fill="x", padx=20, pady=6)

    # --- File picker ---
    ctk.CTkLabel(frame, text="1.  Word file (.docx)",
                 font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=16, pady=(14, 2))
    file_row = ctk.CTkFrame(frame, fg_color="transparent")
    file_row.pack(fill="x", padx=16)
    file_label = ctk.CTkLabel(file_row, text="No file selected", anchor="w",
                              text_color="gray70")
    file_label.pack(side="left", fill="x", expand=True)

    def pick_file():
        p = filedialog.askopenfilename(
            title="Select the medical report",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
        if p:
            state["path"] = p
            file_label.configure(text=os.path.basename(p), text_color="white")

    ctk.CTkButton(file_row, text="Browse…", width=110, command=pick_file).pack(side="right")

    # --- Lab date ---
    ctk.CTkLabel(frame, text="2.  New lab date  (Clinical Chemistry & Hematology)",
                 font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=16, pady=(18, 2))
    ctk.CTkLabel(frame, text="Format: M/D/YYYY", font=ctk.CTkFont(size=11),
                 text_color="gray60").pack(anchor="w", padx=16)
    date_var = ctk.StringVar(value=fmt_labdate(datetime.now()))
    ctk.CTkEntry(frame, textvariable=date_var, width=200).pack(anchor="w", padx=16, pady=(2, 18))

    # --- Run button ---
    status = ctk.CTkLabel(app, text="", font=ctk.CTkFont(size=12), wraplength=560, justify="left")
    status.pack(fill="x", padx=24, pady=(8, 4))

    def run():
        if not state["path"]:
            messagebox.showwarning("Missing file", "Please select a .docx file first.")
            return
        try:
            out, summary = process_report(state["path"], date_var.get())
        except Exception as e:
            status.configure(text=f"Error: {e}", text_color="#ff6b6b")
            return
        msg = (f"Done.\n"
               f"Saved:  {os.path.basename(out)}\n"
               f"Date/time set to:  {summary['now']}  (signature {summary['signature']})\n"
               f"Lab date:  {summary['lab_date']}   "
               f"Tables rolled:  {', '.join(summary['lab_tables_rolled']) or 'none'}")
        status.configure(text=msg, text_color="#4ade80")
        messagebox.showinfo("Success", f"Saved:\n{out}")

    ctk.CTkButton(app, text="Generate Updated Report", height=44,
                  font=ctk.CTkFont(size=15, weight="bold"), command=run).pack(pady=14)

    app.mainloop()


if __name__ == "__main__":
    launch_gui()
